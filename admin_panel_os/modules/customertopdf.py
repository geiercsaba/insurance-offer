import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import modules.doc_util as af
import os

D_PATH = os.environ.get('D_PATH')
NEW_PATH = os.environ.get('NEW_PATH')
D_PATH = "files"
NEW_PATH = "files"

class CreateDOC:
    def __init__(self, customer):
        self.customer = customer
        self.doc = None
        self.new_doc_path = D_PATH
        self.create_template()
        self.document_data()

        self.doc.save(self.new_doc_path+".docx")
        os.startfile(self.new_doc_path+".docx")


    def create_template(self):
        for group in self.customer.groups:
            if group.cmr:
                self.customer.cmr = True
            if group.baf:
                self.customer.baf = True
        if self.customer.cmr and self.customer.baf:
            self.new_doc_path = f"{NEW_PATH}/{self.customer.company_name}_CMR_BAF_Ajánlat_{self.customer.broker}"
            self.doc = docx.Document(f'{D_PATH}/doc_template/cmrbaf_t.docx')
        elif self.customer.cmr:
            self.new_doc_path = f"{NEW_PATH}/{self.customer.company_name}_CMR_Ajánlat_{self.customer.broker}"
            self.doc = docx.Document(f'{D_PATH}/doc_template/cmr_t.docx')
        else:
            self.new_doc_path = f"{NEW_PATH}/{self.customer.company_name}_BAF_Ajánlat_{self.customer.broker}"
            self.doc = docx.Document(f'{D_PATH}/doc_template/baf_t.docx')

    def document_data(self):
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            if '@company_data' in text:
                af.set_company_data(paragraph, self.customer)
            elif '@transported_goods' in text:
                af.set_transported_goods(paragraph, self.customer)
            elif '@excluded_goods' in text:
                af.set_excluded_goods(paragraph, self.customer)
            elif "@number_of_trucks" in text:
                af.set_number_of_truck(paragraph, self.customer)
            elif "@territorial_scope" in text:
                af.set_territorial_scope(paragraph, self.customer)
            elif "@payment_frequency" in text:
                af.set_payment_frequency(paragraph, self.customer)
            elif "@limit" in text:
                af.set_limit(paragraph, self.customer)
            elif "@fee_truck" in text:
                af.set_fee_per_truck(paragraph, self.customer)
            elif "@all_fee" in text:
                af.set_all_fee(paragraph, self.customer)
            elif "@additional" in text:
                af.set_additional(paragraph, self.customer)